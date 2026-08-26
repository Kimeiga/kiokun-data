#!/usr/bin/env python3
"""Build the Kiokun language-course research report as a polished DOCX."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "report-source.md"
OUTPUT = ROOT / "kiokun-language-course-research-2026-08-25.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6770"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GOLD = "8A5A00"
BORDER = "C8D2DC"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError(f"Column widths must total {CONTENT_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def style_callout_paragraph(paragraph, fill=CALLOUT, border=BLUE_GRAY) -> None:
    """Apply a Word-native shaded callout without using a layout table."""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "6")
        node.set(qn("w:color"), border)
        borders.append(node)
    p_pr.append(borders)
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)


def add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_hyperlink(paragraph, text: str, url: str, bold=False, italic=False) -> None:
    rel = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*.*?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text: str, *, base_bold=False, base_italic=False) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            run.bold = base_bold
            run.italic = base_italic
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.italic = base_italic
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            run.bold = base_bold
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Menlo"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Menlo")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url, base_bold, base_italic)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = base_bold
        run.italic = base_italic


def add_numbering_definitions(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    specs = [(100, 100, "bullet", "•"), (101, 101, "decimal", "%1.")]
    for abstract_id, num_id, fmt, marker in specs:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), marker)
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)
    return 100, 101


def add_number_instance(doc: Document, abstract_id: int, num_id: int, start_at: int = 1) -> int:
    numbering = doc.part.numbering_part.element
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), str(start_at))
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_number(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.left_indent = Inches(0)
    normal.paragraph_format.first_line_indent = Inches(0)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.left_indent = Inches(0)
        style.paragraph_format.first_line_indent = Inches(0)

    # Explicit preset-controlled list geometry and rhythm.
    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    # Running header/footer for pages after the cover.
    header = section.header
    p = header.paragraphs[0]
    p.text = "KIOKUN  /  LANGUAGE COURSE RESEARCH"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("Research cutoff: 25 Aug 2026   •   ")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(fp, "PAGE")

    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared for Kiokun product, curriculum, linguistics, and engineering")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(54)
    r = p.add_run("RESEARCH & PRODUCT BLUEPRINT")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Building an Effective\nKiokun Language Course")
    r.font.name = "Calibri"
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("Japanese course landscape, evidence review, and a reusable architecture for Mandarin, Cantonese, and Korean")
    r.font.name = "Calibri"
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    style_callout_paragraph(p)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("DECISION\n")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r = p.add_run("Build Japanese first. Combine Can-do sequencing, precise explanation, cumulative input and retrieval, open production with repair, and delayed transfer assessment in one auditable learning graph.")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(INK)

    doc.add_paragraph()
    for label, value in (
        ("Audience", "Kiokun product, curriculum, linguistics, and engineering teams"),
        ("Research cutoff", "August 25, 2026"),
        ("Market set", "27 Japanese courses, curricula, and specialist systems"),
        ("Primary lens", "Delayed, novel, minimally scaffolded language use"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{label}: ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(INK)
        r = p.add_run(value)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_page_break()


def choose_widths(headers: list[str]) -> list[int]:
    n = len(headers)
    lowered = [h.lower() for h in headers]
    if n == 2:
        return [2700, 6660]
    if n == 3:
        return [2800, 1660, 4900]
    if n == 4 and "resource" in lowered[0]:
        return [2300, 3240, 1680, 2140]
    if n == 4:
        return [2050, 1450, 4100, 1760]
    base = CONTENT_DXA // n
    widths = [base] * n
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    headers = rows[0]
    data = rows[2:] if len(rows) > 1 and all(set(c.strip()) <= set(":-") for c in rows[1]) else rows[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.allow_autofit = False
    set_table_geometry(table, choose_widths(headers))
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BLUE_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_inline(p, value.strip(), base_bold=True)
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(INK)
    for ridx, row_values in enumerate(data):
        row = table.add_row()
        if ridx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, "F8FAFC")
        for idx, value in enumerate(row_values):
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            add_inline(p, value.strip())
            for run in p.runs:
                run.font.size = Pt(8.6)
            if re.fullmatch(r"[A-D][+-]?|High|Medium|Low|\d+", value.strip()):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def parse_markdown(doc: Document, text: str, bullet_id: int, number_id: int) -> None:
    lines = text.splitlines()
    start = next((i + 1 for i, line in enumerate(lines) if line.strip() == "---"), 0)
    lines = lines[start:]
    i = 0
    current_order_num = None
    next_num_id = 200
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            current_order_num = None
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [[c.strip() for c in row.strip("|").split("|")] for row in table_lines]
            add_markdown_table(doc, rows)
            continue
        if line.startswith("### "):
            current_order_num = None
            p = doc.add_paragraph(style="Heading 2")
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.first_line_indent = Inches(0)
            add_inline(p, line[4:])
            i += 1
            continue
        if line.startswith("## "):
            current_order_num = None
            title = line[3:]
            p = doc.add_paragraph(style="Heading 1")
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.first_line_indent = Inches(0)
            add_inline(p, title)
            i += 1
            continue
        if line.startswith("# "):
            current_order_num = None
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[2:])
            i += 1
            continue
        if line.startswith("> "):
            current_order_num = None
            p = doc.add_paragraph()
            style_callout_paragraph(p)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(5)
            add_inline(p, line[2:], base_bold=True)
            i += 1
            continue
        unordered = re.match(r"^-\s+(.*)", line)
        ordered = re.match(r"^\d+\.\s+(.*)", line)
        if unordered or ordered:
            p = doc.add_paragraph()
            if ordered:
                if current_order_num is None:
                    current_order_num = add_number_instance(doc, number_id, next_num_id, 1)
                    next_num_id += 1
                apply_number(p, current_order_num)
            else:
                current_order_num = None
                apply_number(p, bullet_id)
            add_inline(p, (unordered or ordered).group(1))
            i += 1
            continue
        if line == "---":
            current_order_num = None
            i += 1
            continue
        current_order_num = None
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.first_line_indent = Inches(0)
        add_inline(p, line)
        i += 1


def add_source_note(doc: Document) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Source and confidence note")
    p = doc.add_paragraph()
    add_inline(
        p,
        "Claims are linked to official product pages, official standards, or scholarly records near the relevant text. A separate internal claim-to-source ledger accompanies this report in the workspace. Product pages establish scope and documented features; they do not establish causal effectiveness. No complete head-to-head trial of the 27-resource field was found.",
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Prepared 25 August 2026  •  Kiokun")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(MUTED)


def build() -> None:
    doc = Document()
    style_document(doc)
    bullet_id, number_id = add_numbering_definitions(doc)
    doc.core_properties.title = "Building an Effective Kiokun Language Course"
    doc.core_properties.subject = "Japanese course research and multilingual course architecture"
    doc.core_properties.author = "Kiokun Research"
    doc.core_properties.keywords = "Japanese, language learning, Kiokun, Mandarin, Cantonese, Korean"
    add_cover(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"), bullet_id, number_id)
    add_source_note(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
